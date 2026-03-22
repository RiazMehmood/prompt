"""DOCXExportService: export generated documents to DOCX using python-docx.

Supports RTL sections for Urdu/Sindhi with proper font embedding.
"""
import io
from typing import Any, Dict, Optional

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

logger = structlog.get_logger(__name__)

_URDU_FONT = "Jameel Noori Nastaleeq"
_SINDHI_FONT = "Jameel Noori Nastaleeq"
_DEFAULT_FONT = "Times New Roman"


def _set_rtl_paragraph(paragraph: Any) -> None:
    """Set a paragraph to RTL direction via direct XML manipulation."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    # Set text direction
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)


class DOCXExportService:
    """Generate DOCX from rendered document content."""

    def export(
        self,
        content: str,
        output_language: str = "english",
        title: str = "Document",
        formatting_rules: Optional[Dict[str, Any]] = None,
    ) -> bytes:
        """Render content to DOCX bytes.

        RTL languages (Urdu, Sindhi) set section bidi flag + right-aligned paragraphs.
        """
        is_rtl = output_language in ("urdu", "sindhi")
        rules = formatting_rules or {}
        font_name = _URDU_FONT if is_rtl else rules.get("font", _DEFAULT_FONT)
        font_size = int(rules.get("font_size", 12))

        doc = Document()

        # Set section RTL flag for Urdu/Sindhi
        if is_rtl:
            section = doc.sections[0]
            sectPr = section._sectPr
            bidi = OxmlElement("w:bidi")
            sectPr.append(bidi)

        # Add title paragraph
        title_para = doc.add_heading(title, level=1)
        if is_rtl:
            _set_rtl_paragraph(title_para)
            title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in title_para.runs:
                run.font.name = font_name

        # Add body content
        for line in content.split("\n"):
            if not line.strip():
                doc.add_paragraph("")
                continue

            para = doc.add_paragraph()
            run = para.add_run(line.strip())
            run.font.name = font_name
            run.font.size = Pt(font_size)

            if is_rtl:
                _set_rtl_paragraph(para)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
